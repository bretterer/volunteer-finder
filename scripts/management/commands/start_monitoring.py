"""
Monitor folders for new resumes and opportunities simultaneously.
Usage: python manage.py start_monitoring
"""

from django.core.management.base import BaseCommand
from django.utils import timezone
from pathlib import Path
import time
from datetime import datetime, timedelta

from accounts.models import User, VolunteerProfile
from resumes.models import Resume
from resumes.services import ResumeScoringService
from opportunities.models import Opportunity
import re


class Command(BaseCommand):
    help = 'Monitor folders for new resumes and opportunities'

    def add_arguments(self, parser):
        parser.add_argument(
            '--resume-folder',
            type=str,
            default='resume_uploads',
            help='Folder to monitor for new resumes'
        )
        parser.add_argument(
            '--opportunity-folder',
            type=str,
            default='opportunity_uploads',
            help='Folder to monitor for new opportunities'
        )
        parser.add_argument(
            '--interval',
            type=int,
            default=30,
            help='Check interval in seconds (default: 30)'
        )
        parser.add_argument(
            '--auto-score',
            action='store_true',
            help='Automatically score new resumes and opportunities'
        )
        parser.add_argument(
            '--once',
            action='store_true',
            help='Check once and exit (don\'t loop)'
        )

    def handle(self, *args, **options):
        resume_folder = Path(options['resume_folder'])
        opportunity_folder = Path(options['opportunity_folder'])
        interval = options['interval']
        auto_score = options['auto_score']
        once = options['once']

        # Verify folders exist
        if not resume_folder.exists():
            self.stdout.write(self.style.ERROR(f"❌ Resume folder not found: {resume_folder}"))
            return

        if not opportunity_folder.exists():
            self.stdout.write(self.style.ERROR(f"❌ Opportunity folder not found: {opportunity_folder}"))
            return

        # Display header
        self.stdout.write(self.style.SUCCESS('=' * 80))
        self.stdout.write(self.style.SUCCESS('👀 UNIFIED FILE MONITOR'))
        self.stdout.write(self.style.SUCCESS('=' * 80))
        self.stdout.write(f"\n📁 Resume folder: {resume_folder.absolute()}")
        self.stdout.write(f"📁 Opportunity folder: {opportunity_folder.absolute()}")
        self.stdout.write(f"⏱️  Interval: {interval} seconds ({interval / 60:.1f} minutes)")
        self.stdout.write(f"🤖 Auto-score: {'Yes' if auto_score else 'No'}")
        self.stdout.write(f"🔄 Mode: {'Single check' if once else 'Continuous'}\n")

        if not once:
            self.stdout.write(self.style.WARNING("Press Ctrl+C to stop\n"))

        # Main loop
        try:
            while True:
                self.stdout.write('\n' + '=' * 80)
                self.stdout.write(f'🔄 FILE CHECK - {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
                self.stdout.write('=' * 80)

                # Check both folders
                self.check_resumes(resume_folder, auto_score)
                self.stdout.write('')  # Blank line
                self.check_opportunities(opportunity_folder, auto_score)

                if once:
                    break

                # Wait for next check
                next_check = datetime.now().timestamp() + interval
                self.stdout.write(f"\n⏳ Waiting {interval} seconds until next check...")
                self.stdout.write(
                    f"   Next check at: {datetime.fromtimestamp(next_check).strftime('%Y-%m-%d %H:%M:%S')}")

                time.sleep(interval)

        except KeyboardInterrupt:
            self.stdout.write(self.style.SUCCESS('\n\n🛑 MONITORING STOPPED'))

    def check_resumes(self, folder: Path, auto_score: bool):
        """Check folder for new resume files."""
        self.stdout.write('\n📄 CHECKING RESUMES...')

        # Supported extensions
        extensions = ['.pdf', '.docx', '.txt']

        # Find all resume files
        resume_files = []
        for ext in extensions:
            resume_files.extend(folder.glob(f'*{ext}'))

        self.stdout.write(f"   Found {len(resume_files)} files in {folder.name}")

        # Check which are new
        new_files = []
        for file_path in resume_files:
            filename = file_path.name
            exists = Resume.objects.filter(original_filename=filename).exists()
            if not exists:
                new_files.append(file_path)

        if not new_files:
            self.stdout.write("   ✅ No new resume files")
            return

        self.stdout.write(f"   🆕 Found {len(new_files)} new resume(s):")
        for file_path in new_files:
            self.stdout.write(f"      • {file_path.name}")

        # Process new files
        added_resumes = []
        for file_path in new_files:
            resume = self.add_resume_to_database(file_path)
            if resume:
                added_resumes.append(resume)

        self.stdout.write(f"   ✅ Added {len(added_resumes)} resume(s) to database")

        # Auto-score if requested
        if auto_score and added_resumes:
            self.stdout.write("\n   🤖 SCORING NEW RESUMES AGAINST ALL OPPORTUNITIES...")
            service = ResumeScoringService()

            for resume in added_resumes:
                self.stdout.write(f"      Scoring: {resume.original_filename}")
                try:
                    scores = service.score_resume_for_all_opportunities(resume)
                    self.stdout.write(f"      ✅ Created {len(scores)} scores")
                except Exception as e:
                    self.stdout.write(self.style.ERROR(f"      ❌ Error: {e}"))

    def check_opportunities(self, folder: Path, auto_score: bool):
        """Check folder for new opportunity files."""
        self.stdout.write('💼 CHECKING OPPORTUNITIES...')

        # Supported extensions
        extensions = ['.pdf', '.txt', '.docx']

        # Find all opportunity files
        opportunity_files = []
        for ext in extensions:
            opportunity_files.extend(folder.glob(f'*{ext}'))

        self.stdout.write(f"   Found {len(opportunity_files)} files in {folder.name}")

        if not opportunity_files:
            self.stdout.write("   ✅ No opportunity files")
            return

        # Parse and check each file
        new_opportunities = []

        for file_path in opportunity_files:
            filename = file_path.name

            # Check if this exact filename has been processed
            exists = Opportunity.objects.filter(source_filename=filename).exists()

            if exists:
                continue  # Skip already processed files

            # Extract and parse
            extracted_text = self.extract_text_opportunity(file_path)
            if not extracted_text:
                self.stdout.write(self.style.WARNING(f"      ⚠️  No text extracted from {filename}"))
                continue

            opp_data = self.parse_opportunity_text(extracted_text, filename)
            opp_data['source_filename'] = filename  # Track filename

            new_opportunities.append((file_path, opp_data))

        if not new_opportunities:
            self.stdout.write("   ✅ No new opportunity files")
            return

        self.stdout.write(f"   🆕 Found {len(new_opportunities)} new opportunity/ies:")
        for file_path, opp_data in new_opportunities:
            self.stdout.write(f"      • {file_path.name} → {opp_data['title'][:50]}")

        # Process new opportunities
        added_opportunities = []
        for file_path, opp_data in new_opportunities:
            opportunity = self.create_opportunity_from_data(opp_data)
            if opportunity:
                added_opportunities.append(opportunity)

        self.stdout.write(f"   ✅ Added {len(added_opportunities)} opportunity/ies to database")

        # Auto-score if requested
        if auto_score and added_opportunities:
            self.stdout.write("\n   🤖 SCORING ALL RESUMES AGAINST NEW OPPORTUNITIES...")
            service = ResumeScoringService()

            # Get all active resumes
            resumes = Resume.objects.filter(processed=True)
            total_resumes = resumes.count()

            if total_resumes == 0:
                self.stdout.write("      ⚠️  No resumes in database to score")
                return

            self.stdout.write(f"      Found {total_resumes} resumes to score")

            for opportunity in added_opportunities:
                self.stdout.write(f"\n      Scoring all resumes for: {opportunity.title[:50]}")
                scores_created = 0

                for idx, resume in enumerate(resumes, 1):
                    try:
                        # Score this resume for this opportunity
                        score = service.score_resume_for_opportunity(resume, opportunity, force=False)
                        if score:
                            scores_created += 1

                        # Progress indicator every 10 resumes
                        if idx % 10 == 0:
                            self.stdout.write(f"         Progress: {idx}/{total_resumes} resumes scored")

                    except Exception as e:
                        self.stdout.write(self.style.ERROR(f"         ❌ Error scoring resume {resume.id}: {e}"))

                self.stdout.write(f"      ✅ Created {scores_created} scores for {opportunity.title[:50]}")

    # ==================== RESUME PROCESSING ====================

    def add_resume_to_database(self, file_path: Path) -> Resume:
        """Add a resume file to the database."""
        try:
            from django.core.files import File

            filename = file_path.name
            user = self.get_or_create_volunteer(filename)
            extracted_text = self.extract_text_resume(file_path)

            with open(file_path, 'rb') as f:
                resume = Resume(
                    user=user,
                    original_filename=filename,
                    file_size=file_path.stat().st_size,
                    extracted_text=extracted_text,
                    processed=True
                )
                resume.file.save(filename, File(f), save=False)
                resume.save()

            return resume

        except Exception as e:
            self.stdout.write(self.style.ERROR(f"      ❌ Error adding {file_path.name}: {e}"))
            return None

    def extract_text_resume(self, file_path: Path) -> str:
        """Extract text from resume file."""
        ext = file_path.suffix.lower()

        try:
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()

            elif ext == '.pdf':
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ''
                    for page in reader.pages:
                        text += page.extract_text() + '\n'
                    return text

            elif ext == '.docx':
                import docx
                doc = docx.Document(file_path)
                text = '\n'.join([para.text for para in doc.paragraphs])
                return text

            else:
                return ''

        except Exception as e:
            self.stdout.write(self.style.WARNING(f"      ⚠️  Text extraction failed: {e}"))
            return ''

    def get_or_create_volunteer(self, filename: str) -> User:
        """Get or create volunteer user from filename."""
        base_name = Path(filename).stem.replace('_Resume', '').replace('_resume', '').replace('_', ' ')
        username = base_name.lower().replace(' ', '_')[:30]

        user = User.objects.filter(username=username).first()
        if user:
            return user

        user = User.objects.create_user(
            username=username,
            email=f'{username}@volunteer.com',
            user_type='volunteer'
        )

        VolunteerProfile.objects.create(user=user)
        return user

    # ==================== OPPORTUNITY PROCESSING ====================

    def create_opportunity_from_data(self, opp_data: dict) -> Opportunity:
        """Create opportunity from parsed data."""
        try:
            org_user = self.get_or_create_organization(opp_data['organization_name'])

            opportunity = Opportunity.objects.create(
                organization=org_user,
                title=opp_data['title'],
                description=opp_data['description'],
                location=opp_data['location'],
                required_skills=opp_data['required_skills'],
                hours_required=opp_data['hours_required'],
                spots_available=opp_data['spots_available'],
                start_date=opp_data['start_date'],
                end_date=opp_data['end_date'],
                source_filename=opp_data.get('source_filename', ''),
                status='active'
            )

            return opportunity

        except Exception as e:
            self.stdout.write(self.style.ERROR(f"      ❌ Error creating opportunity: {e}"))
            return None

    def extract_text_opportunity(self, file_path: Path) -> str:
        """Extract text from opportunity file."""
        ext = file_path.suffix.lower()

        try:
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()

            elif ext == '.pdf':
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ''
                    for page in reader.pages:
                        text += page.extract_text() + '\n'
                    return text

            elif ext == '.docx':
                import docx
                doc = docx.Document(file_path)
                text = '\n'.join([para.text for para in doc.paragraphs])
                return text

            else:
                return ''

        except Exception as e:
            self.stdout.write(self.style.WARNING(f"      ⚠️  Text extraction failed for {file_path.name}: {e}"))
            return ''

    def parse_opportunity_text(self, text: str, filename: str) -> dict:
        """Parse opportunity details from extracted text."""
        # Default dates: start today, end in 6 months
        default_start = timezone.now().date()
        default_end = (timezone.now() + timedelta(days=180)).date()

        data = {
            'title': filename.replace('.pdf', '').replace('.txt', '').replace('.docx', '').replace('_', ' '),
            'description': text,
            'location': 'Not specified',
            'organization_name': 'System Admin',
            'required_skills': [],
            'hours_required': 5,
            'spots_available': 1,
            'start_date': default_start,
            'end_date': default_end
        }

        # Extract POSITION
        position_match = re.search(r'POSITION:\s*(.+?)(?:\n|DEPARTMENT:)', text, re.IGNORECASE)
        if position_match:
            data['title'] = position_match.group(1).strip()

        # Extract DEPARTMENT
        dept_match = re.search(r'DEPARTMENT:\s*(.+?)(?:\n|Volunteers)', text, re.IGNORECASE)
        if dept_match:
            data['location'] = dept_match.group(1).strip()

        # Extract Volunteers Needed
        spots_match = re.search(r'Volunteers Needed:\s*(\d+)', text, re.IGNORECASE)
        if spots_match:
            data['spots_available'] = int(spots_match.group(1))

        # Extract organization leaders
        org_match = re.search(r'ORGANIZATIONAL LEADERS?:\s*(.+?)(?:\n\n|REQUIREMENTS|$)', text,
                              re.IGNORECASE | re.DOTALL)
        if org_match:
            leaders = org_match.group(1).strip()
            first_leader = re.search(r'[-•]\s*(?:Prof\.|Dr\.)?\s*(.+?)(?:\n|$)', leaders)
            if first_leader:
                data['organization_name'] = first_leader.group(1).strip()

        # Extract hours
        hours_match = re.search(r'(\d+)\s*hours?\s*(?:per week|weekly)', text, re.IGNORECASE)
        if hours_match:
            data['hours_required'] = int(hours_match.group(1))

        return data

    def get_or_create_organization(self, org_name: str) -> User:
        """Get or create organization user."""
        username = org_name.lower().replace(' ', '_').replace('.', '')[:30]

        user = User.objects.filter(username=username).first()
        if user:
            return user

        user = User.objects.create_user(
            username=username,
            email=f'{username}@university.edu',
            user_type='organization'
        )

        return user