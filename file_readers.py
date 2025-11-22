"""
File reader module for processing resumes in various formats.
"""

from pathlib import Path
from typing import Dict, Optional, List
import logging

import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)


class FileReadError(Exception):
    """Custom exception for file reading errors"""
    pass


class UnsupportedFormatError(FileReadError):
    """Exception for unsupported file formats"""
    pass


class ResumeReader:
    """Interface for reading resume files in multiple formats"""

    SUPPORTED_FORMATS: List[str] = ['.pdf', '.docx', '.txt']

    @staticmethod
    def read_file(file_path: str | Path) -> str:
        """
        Read a file and return its text content.

        Args:
            file_path: Path to the resume file

        Returns:
            str: Extracted text content

        Raises:
            FileNotFoundError: If the file doesn't exist
            UnsupportedFormatError: If the file format is not supported
            FileReadError: If there's an error reading the file
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()
        if extension not in ResumeReader.SUPPORTED_FORMATS:
            raise UnsupportedFormatError(
                f"Unsupported format: {extension}. "
                f"Supported formats: {', '.join(ResumeReader.SUPPORTED_FORMATS)}"
            )

        try:
            if extension == '.txt':
                return ResumeReader._read_txt(file_path)
            elif extension == '.docx':
                return ResumeReader._read_docx(file_path)
            elif extension == '.pdf':
                return ResumeReader._read_pdf(file_path)
        except (FileNotFoundError, UnsupportedFormatError):
            raise
        except Exception as e:
            raise FileReadError(
                f"Error reading {file_path.name}: {str(e)}"
            ) from e

    @staticmethod
    def _read_txt(file_path: Path) -> str:
        """
        Read text file content.

        Args:
            file_path: Path to the text file

        Returns:
            str: File content

        Raises:
            FileReadError: If there's an error reading the file
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as txt_file:
                text = txt_file.read()

            if not text.strip():
                logger.warning(f"Text file {file_path.name} is empty")

            logger.debug(f"Read TXT: {file_path.name} ({len(text)} characters)")
            return text

        except Exception as e:
            raise FileReadError(
                f"Failed to read text file {file_path.name}: {str(e)}"
            ) from e

    @staticmethod
    def _read_docx(file_path: Path) -> str:
        """
        Read DOCX file content.

        Args:
            file_path: Path to the DOCX file

        Returns:
            str: Extracted text content

        Raises:
            FileReadError: If there's an error reading the file
        """
        try:
            doc = Document(file_path)
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            text = '\n'.join(text_parts)

            if not text.strip():
                logger.warning(f"DOCX file {file_path.name} appears to be empty")

            logger.debug(f"Read DOCX: {file_path.name} ({len(text)} characters)")
            return text

        except Exception as e:
            raise FileReadError(
                f"Failed to read DOCX file {file_path.name}: {str(e)}"
            ) from e

    @staticmethod
    def _read_pdf(file_path: Path) -> str:
        """
        Read PDF file content.

        Args:
            file_path: Path to the PDF file

        Returns:
            str: Extracted text content

        Raises:
            FileReadError: If there's an error reading the file
        """
        try:
            text_parts = []

            with open(file_path, 'rb') as file_pdf:
                pdf_reader = PyPDF2.PdfReader(file_pdf)
                num_pages = len(pdf_reader.pages)

                if num_pages == 0:
                    raise FileReadError(f"PDF file {file_path.name} has no pages")

                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()

                    if text.strip():
                        text_parts.append(text)

            full_text = '\n'.join(text_parts)

            if not full_text.strip():
                logger.warning(
                    f"PDF file {file_path.name} has {num_pages} pages but no extractable text"
                )

            logger.debug(
                f"Read PDF: {file_path.name} ({len(full_text)} characters from {num_pages} pages)"
            )
            return full_text

        except PyPDF2.errors.PdfReadError as e:
            raise FileReadError(
                f"Invalid or corrupted PDF file {file_path.name}: {str(e)}"
            ) from e
        except Exception as e:
            raise FileReadError(
                f"Failed to read PDF file {file_path.name}: {str(e)}"
            ) from e

    @staticmethod
    def get_file_info(file_path: str | Path) -> Dict[str, str | bool | int]:
        """
        Get information about a file without reading its content.

        Args:
            file_path: Path to the file

        Returns:
            Dict containing filename, extension, support status, and size

        Raises:
            FileNotFoundError: If the file doesn't exist
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()

        return {
            "filename": file_path.name,
            "extension": extension,
            "supported": extension in ResumeReader.SUPPORTED_FORMATS,
            "size_bytes": file_path.stat().st_size
        }

    @staticmethod
    def is_supported(file_path: str | Path) -> bool:
        """
        Check if a file format is supported.

        Args:
            file_path: Path to the file

        Returns:
            bool: True if the format is supported, False otherwise
        """
        extension = Path(file_path).suffix.lower()
        return extension in ResumeReader.SUPPORTED_FORMATS

    @staticmethod
    def batch_read_files(file_paths: List[str | Path]) -> Dict[str, str]:
        """
        Read multiple files and return their contents.

        Args:
            file_paths: List of file paths to read

        Returns:
            Dict mapping file names to their text content

        Note:
            Failed files will be logged but won't stop processing of other files
        """
        results = {}

        for file_path in file_paths:
            try:
                text = ResumeReader.read_file(file_path)
                results[Path(file_path).name] = text
            except FileReadError as e:
                logger.error(f"Failed to read {file_path}: {str(e)}")
            except FileNotFoundError as e:
                logger.error(f"File not found: {file_path}")
            except UnsupportedFormatError as e:
                logger.error(f"Unsupported format for {file_path}: {str(e)}")

        return results